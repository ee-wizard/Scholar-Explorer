from docxtpl import DocxTemplate, InlineImage
import json
import os
import base64
import uuid
import glob
import logging
from difflib import get_close_matches
from datetime import datetime
from docx.shared import Mm, Inches
from docx import Document
from io import BytesIO

# Set up logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def find_best_matching_logo(logo_path):
    """
    Finds the best matching logo file even if the name doesn't exactly match.
    
    Args:
        logo_path (str): The logo path from JSON
    
    Returns:
        str: Path to the best matching logo file or empty string if none found
    """
    if os.path.exists(logo_path):
        return logo_path
    
    # Extract filename from path
    filename = os.path.basename(logo_path)
    file_base, file_ext = os.path.splitext(filename)
    
    # Define where to look for logos (current dir and logos dir)
    search_paths = [
        "*.png", "*.jpg", "*.jpeg", "*.gif",
        "logos/*.png", "logos/*.jpg", "logos/*.jpeg", "logos/*.gif",
        "static/logos/*.png", "static/logos/*.jpg", "static/logos/*.jpeg", "static/logos/*.gif",
        "src/static/logos/*.png", "src/static/logos/*.jpg", "src/static/logos/*.jpeg", "src/static/logos/*.gif"
    ]
    
    # Collect all potential logo files
    all_logo_files = []
    for pattern in search_paths:
        all_logo_files.extend(glob.glob(pattern))
    
    if not all_logo_files:
        return ""
    
    # Extract just the filenames without extensions for comparison
    logo_basenames = [os.path.splitext(os.path.basename(f))[0] for f in all_logo_files]
    
    # Find closest match
    matches = get_close_matches(file_base, logo_basenames, n=1, cutoff=0.6)
    
    if not matches:
        return ""
    
    # Get the index of the match
    match_idx = logo_basenames.index(matches[0])
    return all_logo_files[match_idx]

def create_agenda_doc(data, template_path, output_path=None, logo_path=None):
    """
    Core function to create an agenda document from JSON data
    
    Args:
        data: Dictionary or JSON string of agenda data
        template_path: Path to the template DOCX file
        output_path: Path to save the output (generated if None)
        logo_path: Path to logo file, URL, or base64 encoded image from frontend
    
    Returns:
        Path to the generated document
    """
    # Parse JSON if string was provided
    if isinstance(data, str):
        data = json.loads(data)
    
    logger.info(f"Creating document from template: {template_path}")
    logger.info(f"Output will be saved to: {output_path}")
    
    # Make sure output directory exists
    if output_path:
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
    
    # Load the template
    doc = DocxTemplate(template_path)
    
    # Prepare context with more detailed structure
    context = {
        "customer": data.get("customer", ""),
        "date": data.get("date", ""),
        "title": data.get("title", ""),
        "summary": data.get("summary", ""),
        "primaries": data.get("primaries", []),
        "supporting": data.get("supporting", []),
        "agenda_items": data.get("agenda_items", []),
        "attendees": data.get("attendees", []),
        "has_logo": False  # Default to no logo
    }
    
    # Handle logo
    temp_logo_path = None
    if logo_path:
        logger.info(f"Processing logo: {logo_path[:30]}{'...' if len(logo_path) > 30 else ''}")
        
        # Create logos directory if it doesn't exist
        temp_dir = os.path.abspath(os.path.dirname(output_path) if output_path else os.path.join(os.getcwd(), 'temp'))
        os.makedirs(temp_dir, exist_ok=True)
        
        try:
            # Check if it's a base64 encoded image
            if isinstance(logo_path, str) and logo_path.startswith('data:image'):
                try:
                    # Extract the actual base64 data after the comma
                    base64_data = logo_path.split(',')[1]
                    image_data = base64.b64decode(base64_data)
                    
                    # Save the temporary logo file
                    temp_logo_path = os.path.join(temp_dir, f'temp_logo_{uuid.uuid4()}.png')
                    with open(temp_logo_path, 'wb') as f:
                        f.write(image_data)
                    
                    logo_path = temp_logo_path
                    logger.info(f"Converted base64 logo to file: {logo_path}")
                except Exception as e:
                    logger.error(f"Error processing base64 logo: {str(e)}")
                    context["has_logo"] = False
            
            # At this point, logo_path should be a file path
            if os.path.exists(logo_path):
                try:
                    # Check if file is readable
                    with open(logo_path, 'rb') as test_read:
                        _ = test_read.read(1)
                    
                    # Add multiple logo format options to increase template compatibility
                    # The template might be expecting any of these formats
                    context["logo"] = InlineImage(doc, logo_path, width=Mm(50))
                    context["company_logo"] = context["logo"]  # Alternative name
                    context["logo_image"] = context["logo"]    # Another alternative
                    context["has_logo"] = True
                    logger.info(f"Using file path logo: {logo_path}")
                except Exception as e:
                    logger.error(f"Error creating InlineImage from file: {str(e)}")
                    context["has_logo"] = False
            else:
                logger.warning(f"Logo path not valid or file not found: {logo_path}")
        except Exception as e:
            logger.error(f"Unexpected error in logo processing: {str(e)}")
            context["has_logo"] = False
    
    # Inspect the template variables to better understand what's expected
    try:
        # Extract template variables to see what it expects
        template_vars = doc.get_undeclared_template_variables()
        logger.info(f"Template variables: {template_vars}")
        
        # Check if template expects specific logo-related variables
        logo_related_vars = [var for var in template_vars if 'logo' in var.lower()]
        if logo_related_vars and context.get("has_logo"):
            logger.info(f"Logo-related variables in template: {logo_related_vars}")
            # Ensure all logo-related variables are set
            for var in logo_related_vars:
                if var not in context:
                    context[var] = context.get("logo")
    except Exception as e:
        logger.warning(f"Could not inspect template variables: {str(e)}")
    
    # Render the template with the context
    try:
        doc.render(context)
        logger.info("Template rendered successfully")
    except Exception as e:
        logger.error("Error rendering template:")
        logger.exception(e)  # <-- log the full traceback
        logger.error(f"Context keys: {list(context.keys())}")
        logger.error(f"has_logo value: {context.get('has_logo')}")
        logger.info("Check if your DOCX template has a placeholder like {{ logo }} or {{ company_logo }}")
        
        # Try rendering without the logo as a fallback
        try:
            logger.info("Attempting to render template without logo as fallback")
            fallback_context = context.copy()
            # Remove logo-related keys
            for key in list(fallback_context.keys()):
                if 'logo' in key.lower():
                    del fallback_context[key]
            fallback_context["has_logo"] = False
            
            doc = DocxTemplate(template_path)  # Create fresh template
            doc.render(fallback_context)
            logger.info("Template rendered successfully without logo")
        except Exception as fallback_error:
            logger.error("Fallback rendering also failed:")
            logger.exception(fallback_error)
            raise e  # Raise the original error
    
    # Make sure we have an output path
    if not output_path:
        os.makedirs('output', exist_ok=True)
        current_date = datetime.now().strftime("%Y%m%d")
        customer = data.get('customer', 'Customer').replace(' ', '_')
        topic = data.get('topic', data.get('title', 'Meeting')).replace(' ', '_')
        filename = f"{current_date}-{customer}-{topic}Agenda-{uuid.uuid4()}.docx"
        output_path = os.path.join('output', filename)
    
    # Save the document
    try:
        doc.save(output_path)
        logger.info(f"Document saved to: {output_path}")
        
        # Explicitly call post-processing with additional logging
        logger.info("Calling post-processing function...")
        post_process_result = post_process_document(output_path)
        logger.info(f"Post-processing completed: {post_process_result}")
    except Exception as e:
        logger.error(f"Error saving document: {str(e)}")
        raise
    
    # Clean up temporary logo file if it exists
    if temp_logo_path and os.path.exists(temp_logo_path):
        try:
            os.remove(temp_logo_path)
        except Exception as e:
            logger.warning(f"Could not remove temporary logo file: {str(e)}")
    
    return output_path

def post_process_document(docx_path):
    """
    Post-processes the generated DOCX file to:
    1. Remove the first column from agenda items table (if needed)
    2. Adjust column widths for better appearance
    """
    import logging
    logger = logging.getLogger(__name__)
    logger.info(f"Post-processing document: {docx_path}")
    
    try:
        # Open the document
        doc = Document(docx_path)
        
        # Log how many tables exist
        logger.info(f"Document has {len(doc.tables)} tables")
        
        # Find the first table with more than one row (header + content)
        agenda_table = None
        for i, table in enumerate(doc.tables):
            if len(table.rows) > 1:  # More than just header row
                agenda_table = table
                logger.info(f"Using table #{i+1} with {len(table.rows)} rows and {len(table.columns)} columns")
                break
        
        if agenda_table:
            # Count columns before modification
            original_column_count = len(agenda_table.columns)
            logger.info(f"Table originally has {original_column_count} columns")
            
            # Optional: Remove the first column if there are more than 3 columns
            # This is only needed if your template generates an extra column
            if original_column_count > 3:
                logger.info("Removing first column from table")
                try:
                    for row in agenda_table.rows:
                        # Get the XML element for the row
                        xml_row = row._tr
                        # Remove the first cell if it exists
                        if xml_row.tc_lst:
                            xml_row.remove(xml_row.tc_lst[0])
                    logger.info("First column removed successfully")
                except Exception as e:
                    logger.error(f"Error removing first column: {str(e)}")
            
            # Adjust the remaining columns to appropriate widths
            # The number of columns might have changed if we removed the first one
            current_columns = len(agenda_table.columns)
            logger.info(f"Adjusting widths for {current_columns} columns")
            
            if current_columns >= 3:
                agenda_table.columns[0].width = Inches(0.8)   # Time column
                agenda_table.columns[1].width = Inches(1.2)   # Owner column
                agenda_table.columns[2].width = Inches(4.0)   # Topic/Description column
                logger.info("Column widths adjusted successfully")
            elif current_columns == 2:
                agenda_table.columns[0].width = Inches(1.5)   # Time/Owner column
                agenda_table.columns[1].width = Inches(4.5)   # Topic/Description column
                logger.info("Column widths adjusted for 2-column table")
            
            # Save the modified document
            doc.save(docx_path)
            logger.info(f"Document post-processed successfully: {docx_path}")
            return True
        else:
            logger.warning("No suitable table found with more than one row")
            return False
            
    except Exception as e:
        logger.error(f"Error during post-processing: {str(e)}")
        # Don't fail if post-processing has issues
        return False