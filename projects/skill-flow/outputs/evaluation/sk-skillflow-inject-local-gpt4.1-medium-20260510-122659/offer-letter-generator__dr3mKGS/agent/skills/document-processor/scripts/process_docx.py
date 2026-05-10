#!/usr/bin/env python3
"""
Word Document Processing Script
Processa documentos Word: extracao, tracked changes, criacao.
Suporta deteccao de alteracoes OOXML.

Usage:
    python process_docx.py input.docx --mode extract|changes|metadata
"""

import argparse
import json
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path
from typing import Optional
from datetime import datetime


# Namespaces OOXML
NAMESPACES = {
    'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
    'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
    'dc': 'http://purl.org/dc/elements/1.1/',
    'cp': 'http://schemas.openxmlformats.org/package/2006/metadata/core-properties',
}


def check_dependencies() -> dict:
    """Verifica dependencias instaladas."""
    deps = {
        "python-docx": False,
        "lxml": False,
    }

    try:
        import docx
        deps["python-docx"] = True
    except ImportError:
        pass

    try:
        import lxml
        deps["lxml"] = True
    except ImportError:
        pass

    return deps


def extract_with_docx(docx_path: str) -> dict:
    """Extrai texto e estrutura usando python-docx."""
    from docx import Document
    from docx.table import Table

    result = {
        "source": docx_path,
        "paragraphs": [],
        "tables": [],
        "headers": [],
        "footers": [],
        "metadata": {},
    }

    doc = Document(docx_path)

    # Metadata
    core_props = doc.core_properties
    result["metadata"] = {
        "author": core_props.author,
        "created": str(core_props.created) if core_props.created else None,
        "modified": str(core_props.modified) if core_props.modified else None,
        "last_modified_by": core_props.last_modified_by,
        "title": core_props.title,
        "subject": core_props.subject,
        "keywords": core_props.keywords,
        "revision": core_props.revision,
    }

    # Paragrafos
    for i, para in enumerate(doc.paragraphs):
        para_data = {
            "index": i,
            "text": para.text,
            "style": para.style.name if para.style else None,
        }

        # Detectar nivel de heading
        if para.style and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", ""))
                para_data["heading_level"] = level
            except ValueError:
                pass

        result["paragraphs"].append(para_data)

    # Tabelas
    for i, table in enumerate(doc.tables):
        table_data = {
            "index": i,
            "rows": len(table.rows),
            "columns": len(table.columns) if table.rows else 0,
            "data": [],
        }

        for row in table.rows:
            row_data = [cell.text for cell in row.cells]
            table_data["data"].append(row_data)

        result["tables"].append(table_data)

    # Headers e Footers (primeira secao)
    if doc.sections:
        section = doc.sections[0]
        if section.header:
            for para in section.header.paragraphs:
                if para.text.strip():
                    result["headers"].append(para.text)
        if section.footer:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    result["footers"].append(para.text)

    return result


def extract_tracked_changes(docx_path: str) -> dict:
    """
    Extrai tracked changes diretamente do XML OOXML.
    Detecta <w:ins> (insercoes) e <w:del> (delecoes).
    """
    result = {
        "source": docx_path,
        "has_tracked_changes": False,
        "changes": [],
        "summary": {
            "insertions": 0,
            "deletions": 0,
            "format_changes": 0,
        },
    }

    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            # Ler document.xml
            if 'word/document.xml' not in zf.namelist():
                result["error"] = "Arquivo document.xml nao encontrado"
                return result

            xml_content = zf.read('word/document.xml')
            root = ET.fromstring(xml_content)

            # Buscar insercoes <w:ins>
            for ins in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ins'):
                change = {
                    "type": "insertion",
                    "author": ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author'),
                    "date": ins.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date'),
                    "text": "",
                }

                # Extrair texto inserido
                for t in ins.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t'):
                    if t.text:
                        change["text"] += t.text

                if change["text"]:
                    result["changes"].append(change)
                    result["summary"]["insertions"] += 1

            # Buscar delecoes <w:del>
            for del_elem in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}del'):
                change = {
                    "type": "deletion",
                    "author": del_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}author'),
                    "date": del_elem.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}date'),
                    "text": "",
                }

                # Extrair texto deletado
                for t in del_elem.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}delText'):
                    if t.text:
                        change["text"] += t.text

                if change["text"]:
                    result["changes"].append(change)
                    result["summary"]["deletions"] += 1

            # Buscar mudancas de formatacao <w:rPrChange>
            for fmt in root.iter('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rPrChange'):
                result["summary"]["format_changes"] += 1

    except zipfile.BadZipFile:
        result["error"] = "Arquivo DOCX invalido ou corrompido"
    except ET.ParseError as e:
        result["error"] = f"Erro ao parsear XML: {str(e)}"

    result["has_tracked_changes"] = len(result["changes"]) > 0
    return result


def extract_metadata_ooxml(docx_path: str) -> dict:
    """Extrai metadados detalhados do OOXML."""
    result = {
        "source": docx_path,
        "core_properties": {},
        "app_properties": {},
        "custom_properties": {},
    }

    try:
        with zipfile.ZipFile(docx_path, 'r') as zf:
            # Core properties (docProps/core.xml)
            if 'docProps/core.xml' in zf.namelist():
                xml_content = zf.read('docProps/core.xml')
                root = ET.fromstring(xml_content)

                # Mapear campos
                field_map = {
                    '{http://purl.org/dc/elements/1.1/}title': 'title',
                    '{http://purl.org/dc/elements/1.1/}creator': 'creator',
                    '{http://purl.org/dc/elements/1.1/}subject': 'subject',
                    '{http://purl.org/dc/elements/1.1/}description': 'description',
                    '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}keywords': 'keywords',
                    '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy': 'last_modified_by',
                    '{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}revision': 'revision',
                    '{http://purl.org/dc/terms/}created': 'created',
                    '{http://purl.org/dc/terms/}modified': 'modified',
                }

                for elem in root:
                    if elem.tag in field_map:
                        result["core_properties"][field_map[elem.tag]] = elem.text

            # App properties (docProps/app.xml)
            if 'docProps/app.xml' in zf.namelist():
                xml_content = zf.read('docProps/app.xml')
                root = ET.fromstring(xml_content)

                ns = 'http://schemas.openxmlformats.org/officeDocument/2006/extended-properties'
                fields = ['Application', 'AppVersion', 'Company', 'Pages', 'Words', 'Characters']

                for field in fields:
                    elem = root.find(f'{{{ns}}}{field}')
                    if elem is not None and elem.text:
                        result["app_properties"][field.lower()] = elem.text

    except zipfile.BadZipFile:
        result["error"] = "Arquivo DOCX invalido"
    except ET.ParseError as e:
        result["error"] = f"Erro XML: {str(e)}"

    return result


def format_as_markdown(data: dict) -> str:
    """Formata resultado como markdown."""
    lines = []
    lines.append(f"# Documento: {data.get('source', 'unknown')}")
    lines.append("")

    # Metadata
    meta = data.get("metadata", {})
    if meta:
        lines.append("## Metadados")
        lines.append(f"- **Autor**: {meta.get('author', 'N/A')}")
        lines.append(f"- **Titulo**: {meta.get('title', 'N/A')}")
        lines.append(f"- **Criado**: {meta.get('created', 'N/A')}")
        lines.append(f"- **Modificado**: {meta.get('modified', 'N/A')}")
        lines.append("")

    # Conteudo
    lines.append("## Conteudo")
    for para in data.get("paragraphs", []):
        text = para.get("text", "")
        if not text.strip():
            continue

        heading_level = para.get("heading_level")
        if heading_level:
            lines.append(f"{'#' * (heading_level + 2)} {text}")
        else:
            lines.append(text)
        lines.append("")

    # Tabelas
    for table in data.get("tables", []):
        lines.append(f"### Tabela {table.get('index', 0) + 1}")
        table_data = table.get("data", [])
        if table_data:
            header = table_data[0]
            lines.append("| " + " | ".join(str(c or "") for c in header) + " |")
            lines.append("| " + " | ".join("---" for _ in header) + " |")
            for row in table_data[1:]:
                lines.append("| " + " | ".join(str(c or "") for c in row) + " |")
        lines.append("")

    return "\n".join(lines)


def main():
    parser = argparse.ArgumentParser(description="Processa documentos Word")
    parser.add_argument("input", nargs="?", help="Arquivo DOCX de entrada")
    parser.add_argument("--mode", choices=["extract", "changes", "metadata"],
                        default="extract", help="Modo de operacao")
    parser.add_argument("--output", choices=["json", "markdown"], default="json",
                        help="Formato de saida")
    parser.add_argument("--check-deps", action="store_true",
                        help="Verificar dependencias")

    args = parser.parse_args()

    if args.check_deps:
        deps = check_dependencies()
        print(json.dumps(deps, indent=2))
        sys.exit(0 if deps["python-docx"] else 1)

    if not args.input:
        parser.print_help()
        sys.exit(1)

    input_path = Path(args.input)
    if not input_path.exists():
        print(json.dumps({"error": f"Arquivo nao encontrado: {args.input}"}))
        sys.exit(1)

    if args.mode == "extract":
        deps = check_dependencies()
        if not deps["python-docx"]:
            print(json.dumps({
                "error": "python-docx nao instalado",
                "install": "pip install python-docx",
            }))
            sys.exit(1)

        result = extract_with_docx(str(input_path))
        if args.output == "markdown":
            print(format_as_markdown(result))
        else:
            print(json.dumps(result, indent=2, ensure_ascii=False, default=str))

    elif args.mode == "changes":
        result = extract_tracked_changes(str(input_path))
        print(json.dumps(result, indent=2, ensure_ascii=False))

    elif args.mode == "metadata":
        result = extract_metadata_ooxml(str(input_path))
        print(json.dumps(result, indent=2, ensure_ascii=False))


if __name__ == "__main__":
    main()
