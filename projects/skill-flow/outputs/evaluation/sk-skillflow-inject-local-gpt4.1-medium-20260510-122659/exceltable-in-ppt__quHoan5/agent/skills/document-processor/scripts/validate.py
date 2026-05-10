#!/usr/bin/env python3
"""
Document Validation Script
Validacao unificada para documentos PDF, XLSX, DOCX.
Implementa validation-first pattern.

Usage:
    python validate.py input.pdf|xlsx|docx
    python validate.py --check-deps
"""

import argparse
import json
import sys
import subprocess
import shutil
from pathlib import Path


def check_all_dependencies() -> dict:
    """Verifica todas as dependencias do skill."""
    result = {
        "python_packages": {},
        "system_tools": {},
        "overall_status": "ok",
        "missing_required": [],
        "missing_optional": [],
    }

    # Python packages
    packages = {
        "pdfplumber": {"required": True, "purpose": "PDF text extraction"},
        "openpyxl": {"required": True, "purpose": "Excel processing"},
        "python-docx": {"required": True, "purpose": "Word processing", "import_name": "docx"},
        "pandas": {"required": False, "purpose": "Data manipulation"},
        "pytesseract": {"required": False, "purpose": "OCR for scanned PDFs"},
        "pdf2image": {"required": False, "purpose": "PDF to image conversion"},
        "reportlab": {"required": False, "purpose": "PDF creation"},
        "defusedxml": {"required": False, "purpose": "Safe XML parsing"},
    }

    for pkg, info in packages.items():
        import_name = info.get("import_name", pkg.replace("-", "_"))
        try:
            __import__(import_name)
            result["python_packages"][pkg] = {
                "installed": True,
                "required": info["required"],
                "purpose": info["purpose"],
            }
        except ImportError:
            result["python_packages"][pkg] = {
                "installed": False,
                "required": info["required"],
                "purpose": info["purpose"],
            }
            if info["required"]:
                result["missing_required"].append(pkg)
            else:
                result["missing_optional"].append(pkg)

    # System tools
    tools = {
        "pdftotext": {"required": False, "purpose": "PDF text extraction (poppler)", "package": "poppler-utils"},
        "tesseract": {"required": False, "purpose": "OCR engine", "package": "tesseract-ocr"},
        "libreoffice": {"required": False, "purpose": "Excel formula validation", "package": "libreoffice"},
        "pandoc": {"required": False, "purpose": "Document conversion", "package": "pandoc"},
    }

    for tool, info in tools.items():
        path = shutil.which(tool)
        result["system_tools"][tool] = {
            "installed": path is not None,
            "path": path,
            "required": info["required"],
            "purpose": info["purpose"],
            "package": info["package"],
        }
        if not path and info["required"]:
            result["missing_required"].append(f"system:{tool}")
        elif not path:
            result["missing_optional"].append(f"system:{tool}")

    # Overall status
    if result["missing_required"]:
        result["overall_status"] = "missing_required"
    elif result["missing_optional"]:
        result["overall_status"] = "missing_optional"

    # Installation instructions
    if result["missing_required"] or result["missing_optional"]:
        pip_packages = [p for p in result["missing_required"] + result["missing_optional"]
                        if not p.startswith("system:")]
        apt_packages = [result["system_tools"][p.replace("system:", "")]["package"]
                        for p in result["missing_required"] + result["missing_optional"]
                        if p.startswith("system:")]

        result["install_instructions"] = {}
        if pip_packages:
            result["install_instructions"]["pip"] = f"pip install {' '.join(pip_packages)}"
        if apt_packages:
            result["install_instructions"]["apt"] = f"sudo apt install {' '.join(apt_packages)}"

    return result


def detect_file_type(file_path: str) -> str:
    """Detecta tipo de arquivo pela extensao."""
    path = Path(file_path)
    ext = path.suffix.lower()

    type_map = {
        ".pdf": "pdf",
        ".xlsx": "xlsx",
        ".xls": "xlsx",
        ".docx": "docx",
        ".doc": "docx",
    }

    return type_map.get(ext, "unknown")


def validate_pdf(file_path: str) -> dict:
    """Valida arquivo PDF."""
    result = {
        "file": file_path,
        "type": "pdf",
        "status": "valid",
        "errors": [],
        "warnings": [],
        "info": {},
    }

    path = Path(file_path)
    if not path.exists():
        result["status"] = "invalid"
        result["errors"].append({"type": "file_not_found", "message": "Arquivo nao encontrado"})
        return result

    # Verificar se e um PDF valido
    try:
        with open(file_path, "rb") as f:
            header = f.read(8)
            if not header.startswith(b"%PDF"):
                result["status"] = "invalid"
                result["errors"].append({"type": "invalid_pdf", "message": "Arquivo nao e um PDF valido"})
                return result
    except IOError as e:
        result["status"] = "invalid"
        result["errors"].append({"type": "read_error", "message": str(e)})
        return result

    # Tentar abrir com pdfplumber
    try:
        import pdfplumber
        with pdfplumber.open(file_path) as pdf:
            result["info"]["pages"] = len(pdf.pages)
            result["info"]["metadata"] = pdf.metadata

            # Verificar se tem texto extraivel
            text_found = False
            for page in pdf.pages[:3]:  # Primeiras 3 paginas
                text = page.extract_text()
                if text and len(text.strip()) > 50:
                    text_found = True
                    break

            if not text_found:
                result["warnings"].append({
                    "type": "no_text",
                    "message": "PDF parece ser escaneado (sem texto extraivel). Use --ocr para extrair.",
                })

    except Exception as e:
        result["status"] = "warning"
        result["warnings"].append({"type": "parse_warning", "message": str(e)})

    return result


def validate_xlsx(file_path: str) -> dict:
    """Valida arquivo Excel."""
    result = {
        "file": file_path,
        "type": "xlsx",
        "status": "valid",
        "errors": [],
        "warnings": [],
        "info": {},
    }

    path = Path(file_path)
    if not path.exists():
        result["status"] = "invalid"
        result["errors"].append({"type": "file_not_found", "message": "Arquivo nao encontrado"})
        return result

    try:
        import openpyxl

        # Carregar com data_only=True para verificar valores calculados
        wb = openpyxl.load_workbook(file_path, data_only=True)
        result["info"]["sheets"] = wb.sheetnames
        result["info"]["sheet_count"] = len(wb.sheetnames)

        # Verificar erros de formula
        excel_errors = ["#VALUE!", "#REF!", "#DIV/0!", "#NAME?", "#N/A", "#NUM!", "#NULL!"]
        formula_errors = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            for row in ws.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        for err in excel_errors:
                            if cell.value.startswith(err):
                                formula_errors.append({
                                    "sheet": sheet_name,
                                    "cell": cell.coordinate,
                                    "error": cell.value,
                                })

        if formula_errors:
            result["status"] = "errors_found"
            result["errors"] = formula_errors
            result["info"]["total_formula_errors"] = len(formula_errors)

        wb.close()

    except Exception as e:
        result["status"] = "invalid"
        result["errors"].append({"type": "parse_error", "message": str(e)})

    return result


def validate_docx(file_path: str) -> dict:
    """Valida arquivo Word."""
    result = {
        "file": file_path,
        "type": "docx",
        "status": "valid",
        "errors": [],
        "warnings": [],
        "info": {},
    }

    path = Path(file_path)
    if not path.exists():
        result["status"] = "invalid"
        result["errors"].append({"type": "file_not_found", "message": "Arquivo nao encontrado"})
        return result

    # Verificar se e um DOCX valido (arquivo ZIP)
    import zipfile
    if not zipfile.is_zipfile(file_path):
        result["status"] = "invalid"
        result["errors"].append({"type": "invalid_docx", "message": "Arquivo nao e um DOCX valido"})
        return result

    try:
        from docx import Document
        doc = Document(file_path)

        result["info"]["paragraphs"] = len(doc.paragraphs)
        result["info"]["tables"] = len(doc.tables)
        result["info"]["sections"] = len(doc.sections)

        # Verificar tracked changes
        import xml.etree.ElementTree as ET
        with zipfile.ZipFile(file_path, 'r') as zf:
            if 'word/document.xml' in zf.namelist():
                xml_content = zf.read('word/document.xml')
                root = ET.fromstring(xml_content)

                ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
                insertions = len(list(root.iter(f'{ns}ins')))
                deletions = len(list(root.iter(f'{ns}del')))

                if insertions > 0 or deletions > 0:
                    result["info"]["tracked_changes"] = {
                        "insertions": insertions,
                        "deletions": deletions,
                    }
                    result["warnings"].append({
                        "type": "tracked_changes",
                        "message": f"Documento tem {insertions + deletions} alteracoes pendentes",
                    })

    except Exception as e:
        result["status"] = "invalid"
        result["errors"].append({"type": "parse_error", "message": str(e)})

    return result


def main():
    parser = argparse.ArgumentParser(description="Valida documentos PDF, XLSX, DOCX")
    parser.add_argument("input", nargs="?", help="Arquivo para validar")
    parser.add_argument("--check-deps", action="store_true",
                        help="Verificar todas as dependencias")
    parser.add_argument("--strict", action="store_true",
                        help="Modo estrito: warnings viram errors")

    args = parser.parse_args()

    if args.check_deps:
        result = check_all_dependencies()
        print(json.dumps(result, indent=2))
        sys.exit(0 if result["overall_status"] == "ok" else 1)

    if not args.input:
        parser.print_help()
        sys.exit(1)

    file_type = detect_file_type(args.input)

    if file_type == "unknown":
        print(json.dumps({
            "error": f"Tipo de arquivo nao suportado: {args.input}",
            "supported": ["pdf", "xlsx", "xls", "docx", "doc"],
        }))
        sys.exit(1)

    # Validar baseado no tipo
    if file_type == "pdf":
        result = validate_pdf(args.input)
    elif file_type == "xlsx":
        result = validate_xlsx(args.input)
    elif file_type == "docx":
        result = validate_docx(args.input)
    else:
        result = {"error": "Tipo desconhecido"}

    # Modo estrito
    if args.strict and result.get("warnings"):
        result["errors"].extend(result["warnings"])
        result["warnings"] = []
        if result["status"] == "valid":
            result["status"] = "errors_found"

    print(json.dumps(result, indent=2, ensure_ascii=False))

    # Exit code
    if result.get("status") in ["invalid", "errors_found"]:
        sys.exit(1)
    sys.exit(0)


if __name__ == "__main__":
    main()
